import re
import zipfile
from pathlib import Path
from typing import Union

from docx import Document


def extract_docx_text(path: Union[str, Path]) -> str:
    docx_path = Path(path)
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    text = re.sub(r"<[^>]+>", "\n", xml)
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


def save_markdown(markdown: str, path: Union[str, Path]) -> None:
    Path(path).write_text(markdown, encoding="utf-8")


def save_docx_from_markdown(markdown: str, path: Union[str, Path]) -> None:
    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
    document.save(str(path))
